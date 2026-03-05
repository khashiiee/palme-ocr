"""
DOCX Writer — Generate formatted Word documents from dots.ocr structured output.

Uses the layout categories from dots.ocr to apply appropriate formatting.
"""

import json
import re
from typing import Optional

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def parse_ocr_output(raw_output: str) -> Optional[list]:
    """Parse JSON output from dots.ocr — robust against model quirks."""
    text = raw_output.strip()
    text = re.sub(r"^```(?:json)?\s*", "", text)
    text = re.sub(r"\s*```$", "", text)
    text = text.strip()

    # Try direct parse
    try:
        parsed = json.loads(text)
        return _normalize(parsed)
    except json.JSONDecodeError:
        pass

    # Try to find JSON array in mixed text
    match = re.search(r'\[\s*\{.*\}\s*\]', text, re.DOTALL)
    if match:
        try:
            parsed = json.loads(match.group())
            return _normalize(parsed)
        except json.JSONDecodeError:
            pass

    # Fix common JSON issues
    fixed = re.sub(r',\s*([}\]])', r'\1', text)
    try:
        parsed = json.loads(fixed)
        return _normalize(parsed)
    except json.JSONDecodeError:
        pass

    # Regex fallback: extract elements from malformed JSON
    elements = []
    pattern = r'"category"\s*:\s*"([^"]*)"[^}]*?"text"\s*:\s*"((?:[^"\\]|\\.)*)"'
    for m in re.finditer(pattern, text, re.DOTALL):
        cat = m.group(1)
        txt = m.group(2).replace("\\n", "\n").replace('\\"', '"')
        elements.append({"category": cat, "text": txt})

    return elements if elements else None


def _normalize(parsed) -> Optional[list]:
    if isinstance(parsed, list):
        return parsed
    elif isinstance(parsed, dict):
        for key in ["layout_dets", "layout", "elements", "results"]:
            if key in parsed:
                return parsed[key]
        if "text" in parsed:
            return [parsed]
    return None


def clean_html_table(html: str) -> str:
    text = re.sub(r"<tr[^>]*>", "", html)
    text = re.sub(r"</tr>", "\n", text)
    text = re.sub(r"<t[dh][^>]*>", "", text)
    text = re.sub(r"</t[dh]>", "\t", text)
    text = re.sub(r"<[^>]+>", "", text)
    lines = [line.strip() for line in text.split("\n") if line.strip()]
    return "\n".join(lines)


def add_element_to_doc(doc: Document, category: str, text: str):
    """Add a single layout element with appropriate formatting."""
    if not text or category in ("Page-header", "Page-footer"):
        return

    # Clean markdown artifacts
    clean = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    clean = re.sub(r'\*(.+?)\*', r'\1', clean)

    if category == "Title":
        p = doc.add_heading(clean, level=1)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(16)

    elif category == "Section-header":
        # Strip markdown heading markers
        clean = re.sub(r'^#+\s*', '', clean)
        p = doc.add_heading(clean, level=2)
        for run in p.runs:
            run.font.name = "Arial"
            run.font.size = Pt(13)

    elif category == "Picture":
        p = doc.add_paragraph("[Bild / Image]")
        p.italic = True
        for run in p.runs:
            run.font.size = Pt(10)

    elif category == "Caption":
        p = doc.add_paragraph(clean)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.italic = True

    elif category == "Footnote":
        p = doc.add_paragraph(clean)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(8)

    elif category == "Table":
        table_text = clean_html_table(clean) if "<" in clean else clean
        p = doc.add_paragraph(table_text)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    elif category == "Formula":
        p = doc.add_paragraph(clean)
        for run in p.runs:
            run.font.name = "Cambria Math"
            run.font.size = Pt(11)
            run.italic = True

    elif category == "List-item":
        p = doc.add_paragraph(clean, style="List Bullet")
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(11)

    else:
        # Default: Text
        p = doc.add_paragraph(clean)
        for run in p.runs:
            run.font.name = "Courier New"
            run.font.size = Pt(11)


def create_docx_from_pages(pages_raw: list, output_path: str, source_filename: str = ""):
    """Create a formatted Word document from raw OCR outputs (one per page)."""
    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Courier New"
    font.size = Pt(11)

    # Source info header
    if source_filename:
        p = doc.add_paragraph(f"OCR Output: {source_filename}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.italic = True
        doc.add_paragraph("")

    for page_idx, raw_output in enumerate(pages_raw):
        if page_idx > 0:
            doc.add_page_break()

        elements = parse_ocr_output(raw_output)

        if elements and isinstance(elements, list):
            for elem in elements:
                if not isinstance(elem, dict):
                    continue
                category = elem.get("category", "Text").strip()
                text = elem.get("text", "").strip()
                add_element_to_doc(doc, category, text)
        else:
            # Fallback: add raw text but cap it
            fallback = raw_output.strip()[:3000]
            # Try to strip any JSON-looking content
            fallback = re.sub(r'\[\s*\{.*?\}\s*\]', '', fallback, flags=re.DOTALL)
            if fallback.strip():
                doc.add_paragraph(fallback.strip())

    doc.save(output_path)